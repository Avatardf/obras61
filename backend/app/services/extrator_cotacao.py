"""
Extração automática de itens de propostas de cotação.
Suporta: PDF (via Gemini Vision), XLSX (via openpyxl), DOCX (via python-docx).
"""
from __future__ import annotations

import io
import json
import re
from typing import Any

from google import genai
from google.genai import types

from app.config import settings

_client = genai.Client(api_key=settings.gemini_api_key)
_MODEL = "gemini-2.5-flash"

_PROMPT_EXTRACAO = """
Você está analisando uma proposta comercial / cotação de fornecedor de materiais de construção civil.

Extraia TODOS os itens da tabela de preços e retorne um JSON array.
Cada item deve ter os campos:
  - "descricao":      string — descrição completa do material ou serviço
  - "unidade":        string — unidade de medida (ex: "un", "m", "kg", "m²", "pc", "cx", "rl", "bd")
  - "quantidade":     number — quantidade cotada (use 1 se não informada)
  - "preco_unitario": number — preço unitário em reais (apenas número, sem R$ ou pontos de milhar)
  - "observacao":     string ou null — observação do fornecedor para o item, se houver

Regras importantes:
- Retorne APENAS o JSON array, sem markdown, sem texto explicativo
- Preços: converta vírgula decimal para ponto (ex: "1.250,90" → 1250.90)
- Se não houver quantidade na proposta, use 1
- Ignore linhas de totais, subtotais e cabeçalhos
- Se um campo não existir, use null

Exemplo de saída esperada:
[
  {"descricao": "Abraçadeira econômica 2\"", "unidade": "un", "quantidade": 100, "preco_unitario": 1.25, "observacao": null},
  {"descricao": "Cano PVC rígido 100mm", "unidade": "m", "quantidade": 50, "preco_unitario": 18.90, "observacao": "Marca Tigre"}
]
"""


def _parse_json_itens(text: str) -> list[dict]:
    """Extrai e parseia o array JSON da resposta do modelo."""
    # Tenta achar um array JSON na resposta
    match = re.search(r"\[.*\]", text, re.DOTALL)
    if not match:
        return []
    try:
        itens = json.loads(match.group())
        # Normaliza campos
        resultado = []
        for it in itens:
            if not isinstance(it, dict):
                continue
            desc = str(it.get("descricao") or "").strip()
            if not desc:
                continue
            try:
                preco = float(str(it.get("preco_unitario") or 0).replace(",", ".").replace(" ", ""))
            except (ValueError, TypeError):
                preco = 0.0
            try:
                qtd = float(str(it.get("quantidade") or 1).replace(",", "."))
            except (ValueError, TypeError):
                qtd = 1.0
            resultado.append({
                "descricao": desc,
                "unidade": str(it.get("unidade") or "un").strip().lower(),
                "quantidade": qtd,
                "preco_unitario": preco,
                "observacao": it.get("observacao") or None,
            })
        return resultado
    except json.JSONDecodeError:
        return []


async def extrair_de_pdf(pdf_bytes: bytes) -> list[dict[str, Any]]:
    """
    Usa Gemini Vision para extrair itens de uma proposta PDF.
    Funciona com qualquer layout — tabelas, listas, etc.
    """
    parts: list[Any] = [
        types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"),
        types.Part.from_text(text=_PROMPT_EXTRACAO),
    ]
    response = await _client.aio.models.generate_content(
        model=_MODEL,
        contents=parts,
    )
    return _parse_json_itens(response.text)


async def extrair_de_xlsx(xlsx_bytes: bytes) -> list[dict[str, Any]]:
    """
    Extrai itens de planilha Excel.
    Detecta automaticamente colunas por palavras-chave no cabeçalho.
    """
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    ws = wb.active

    COLS_DESC  = {"descricao", "descrição", "material", "item", "produto", "serviço", "servico"}
    COLS_UN    = {"unidade", "un", "und", "unit"}
    COLS_QTD   = {"quantidade", "qtd", "qty", "quant"}
    COLS_PRECO = {"preco", "preço", "valor", "unitario", "unitário", "unit price", "vl unit"}
    COLS_OBS   = {"obs", "observacao", "observação", "nota", "note"}

    # Descobre cabeçalhos
    header_row: int | None = None
    col_map: dict[str, int] = {}

    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=20, values_only=True), start=1):
        hits = 0
        tmp: dict[str, int] = {}
        for col_idx, cell in enumerate(row):
            if cell is None:
                continue
            val = str(cell).lower().strip()
            for key, keywords in [
                ("descricao", COLS_DESC), ("unidade", COLS_UN),
                ("quantidade", COLS_QTD), ("preco_unitario", COLS_PRECO),
                ("observacao", COLS_OBS),
            ]:
                if any(kw in val for kw in keywords):
                    tmp[key] = col_idx
                    hits += 1
        if hits >= 2:
            header_row = row_idx
            col_map = tmp
            break

    if not col_map.get("descricao") and not col_map.get("preco_unitario"):
        # Fallback: assume colunas A=descrição, B=unidade, C=quantidade, D=preço
        col_map = {"descricao": 0, "unidade": 1, "quantidade": 2, "preco_unitario": 3}
        header_row = 1

    resultado: list[dict] = []
    start = (header_row or 1) + 1
    for row in ws.iter_rows(min_row=start, values_only=True):
        desc_val = row[col_map["descricao"]] if "descricao" in col_map and col_map["descricao"] < len(row) else None
        if not desc_val:
            continue
        desc = str(desc_val).strip()
        if not desc or desc.lower() in {"total", "subtotal", ""}:
            continue
        try:
            preco_raw = row[col_map["preco_unitario"]] if "preco_unitario" in col_map and col_map["preco_unitario"] < len(row) else 0
            preco = float(str(preco_raw).replace("R$", "").replace(".", "").replace(",", ".").strip()) if preco_raw else 0.0
        except (ValueError, TypeError):
            preco = 0.0
        try:
            qtd_raw = row[col_map["quantidade"]] if "quantidade" in col_map and col_map["quantidade"] < len(row) else 1
            qtd = float(str(qtd_raw).replace(",", ".")) if qtd_raw else 1.0
        except (ValueError, TypeError):
            qtd = 1.0
        un = str(row[col_map["unidade"]]).strip().lower() if "unidade" in col_map and col_map["unidade"] < len(row) and row[col_map["unidade"]] else "un"
        obs_raw = row[col_map["observacao"]] if "observacao" in col_map and col_map["observacao"] < len(row) else None
        obs = str(obs_raw).strip() if obs_raw else None

        resultado.append({"descricao": desc, "unidade": un, "quantidade": qtd, "preco_unitario": preco, "observacao": obs})

    wb.close()
    return resultado


async def extrair_de_docx(docx_bytes: bytes) -> list[dict[str, Any]]:
    """
    Extrai itens de documento Word (.docx).
    Usa python-docx para ler tabelas, depois Gemini para fallback se necessário.
    """
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    itens: list[dict] = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        # Lê cabeçalho
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        col_desc = col_un = col_qtd = col_preco = col_obs = None
        for i, h in enumerate(header):
            if any(k in h for k in ["descri", "material", "item", "produto"]):
                col_desc = i
            elif any(k in h for k in ["unid", "un"]):
                col_un = i
            elif any(k in h for k in ["quant", "qtd"]):
                col_qtd = i
            elif any(k in h for k in ["preço", "preco", "valor", "unit"]):
                col_preco = i
            elif any(k in h for k in ["obs", "nota"]):
                col_obs = i

        if col_desc is None:
            continue

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= (col_desc or 0):
                continue
            desc = cells[col_desc or 0]
            if not desc:
                continue
            try:
                preco = float(cells[col_preco].replace("R$","").replace(".","").replace(",",".")) if col_preco and col_preco < len(cells) and cells[col_preco] else 0.0
            except (ValueError, TypeError):
                preco = 0.0
            try:
                qtd = float(cells[col_qtd].replace(",",".")) if col_qtd and col_qtd < len(cells) and cells[col_qtd] else 1.0
            except (ValueError, TypeError):
                qtd = 1.0
            un = cells[col_un].lower() if col_un and col_un < len(cells) and cells[col_un] else "un"
            obs = cells[col_obs] if col_obs and col_obs < len(cells) else None
            itens.append({"descricao": desc, "unidade": un, "quantidade": qtd, "preco_unitario": preco, "observacao": obs or None})

    # Se não encontrou tabelas, usa Gemini no texto completo
    if not itens:
        texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if texto:
            parts: list[Any] = [types.Part.from_text(text=f"{_PROMPT_EXTRACAO}\n\nConteúdo do documento:\n{texto}")]
            response = await _client.aio.models.generate_content(model=_MODEL, contents=parts)
            itens = _parse_json_itens(response.text)

    return itens


async def extrair_itens(arquivo_bytes: bytes, mime_type: str) -> list[dict[str, Any]]:
    """
    Dispatcher principal — detecta o tipo e chama o extrator correto.
    mime_type: "application/pdf", "application/vnd.openxmlformats...", "application/msword", etc.
    """
    mt = mime_type.lower()
    if "pdf" in mt:
        return await extrair_de_pdf(arquivo_bytes)
    if "spreadsheet" in mt or "excel" in mt or mt in ("application/vnd.ms-excel",):
        return await extrair_de_xlsx(arquivo_bytes)
    if "wordprocessingml" in mt or "msword" in mt or "docx" in mt:
        return await extrair_de_docx(arquivo_bytes)
    # Fallback: tenta Gemini com qualquer tipo
    return await extrair_de_pdf(arquivo_bytes)
